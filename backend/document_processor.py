import os
import csv
from pathlib import Path
import pdfplumber
from docx import Document as DocxDocument
from datetime import datetime
import pandas as pd
from backend.database import SessionLocal, Document
from backend.ocr_utils import extract_text_from_image
from backend.config import RAW_DIR

def extract_text_from_pdf(path: str) -> str:
    text = ""
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def extract_text_from_docx(path: str) -> str:
    doc = DocxDocument(path)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def extract_text_from_csv_as_docs(file_path: str, source_dataset: str) -> list:
    """
    Reads a CSV file, creates one Document per row.
    Returns list of Document objects (not yet committed).
    """
    df = pd.read_csv(file_path)
    docs = []
    for idx, row in df.iterrows():
        # Combine all text columns into extracted_text
        text_parts = []
        for col in df.columns:
            if col not in ['latitude', 'longitude', 'date posted', 'shape', 'comment_length',
                           'country', 'state', 'verified', 'year_month']:
                val = row[col]
                if pd.notna(val):
                    text_parts.append(str(val))
        full_text = " ".join(text_parts)
        if not full_text.strip():
            continue
        
        # Parse date
        date_val = None
        if 'date posted' in df.columns and pd.notna(row['date posted']):
            try:
                date_val = pd.to_datetime(row['date posted']).to_pydatetime()
            except:
                pass
        
        # year_month for animation
        ym = None
        if date_val:
            ym = date_val.strftime("%Y-%m")
        
        doc = Document(
            filename=f"{Path(file_path).stem}_row{idx}.csv",
            original_path=f"{file_path}:row{idx}",
            extracted_text=full_text,
            text_preview=full_text[:500],
            file_type="csv_row",
            source_dataset=source_dataset,
            doc_date=date_val,
            year_month=ym,
            latitude=float(row['latitude']) if 'latitude' in df.columns and pd.notna(row['latitude']) else None,
            longitude=float(row['longitude']) if 'longitude' in df.columns and pd.notna(row['longitude']) else None,
            incident_shape=str(row['shape']) if 'shape' in df.columns and pd.notna(row['shape']) else None,
            comment_length=int(row['comment_length']) if 'comment_length' in df.columns and pd.notna(row['comment_length']) else None,
            country=str(row['country']) if 'country' in df.columns and pd.notna(row['country']) else None,
            state_code=str(row['state']) if 'state' in df.columns and pd.notna(row['state']) else None,
            verified=bool(row['verified']) if 'verified' in df.columns and pd.notna(row['verified']) else False,
            topic_id=-1,
            anomaly_score=0.0
        )
        docs.append(doc)
    return docs

def extract_text_from_image_file(path: str) -> str:
    return extract_text_from_image(path)

def get_file_date(path: str) -> datetime | None:
    """Attempt to extract date from file metadata (modification time)."""
    try:
        ts = os.path.getmtime(path)
        return datetime.fromtimestamp(ts)
    except:
        return None

def ingest_file(file_path: Path, source_dataset: str = "upload") -> int:
    """Parse file, store text in DB. Returns document ID."""
    ext = file_path.suffix.lower()
    text = ""
    if ext == ".pdf":
        text = extract_text_from_pdf(str(file_path))
    elif ext == ".docx":
        text = extract_text_from_docx(str(file_path))
    elif ext == ".txt":
        text = extract_text_from_txt(str(file_path))
    elif ext == ".csv":
        text = extract_text_from_csv(str(file_path))
    elif ext in (".png", ".jpg", ".jpeg"):
        text = extract_text_from_image_file(str(file_path))
        has_ocr = True
    else:
        return -1

    if not text.strip():
        return -1

    preview = text[:500].replace("\n", " ")
    doc_date = get_file_date(file_path)
    db = SessionLocal()
    doc = Document(
        filename=file_path.name,
        original_path=str(file_path),
        extracted_text=text,
        text_preview=preview,
        file_type=ext[1:],
        source_dataset=source_dataset,
        doc_date=doc_date,
        has_ocr=(ext in (".png", ".jpg", ".jpeg")),
        topic_id=-1,
        anomaly_score=0.0
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    db.close()
    return doc.id

def ingest_all_raw_files():
    """Scan data/raw and ingest any new files."""
    db = SessionLocal()
    existing_paths = {doc.original_path for doc in db.query(Document.original_path).all()}
    db.close()

    for file_path in RAW_DIR.iterdir():
        if file_path.is_file() and file_path.suffix.lower() in {".pdf", ".docx", ".txt", ".csv", ".png", ".jpg", ".jpeg"}:
            if str(file_path) not in existing_paths:
                ingest_file(file_path, source_dataset="raw_folder")
                print(f"Ingested: {file_path.name}")
